from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any

from backend.storage import ensure_upload_root

try:
    from docx import Document  # type: ignore
except Exception:  # pragma: no cover
    Document = None


def _safe_name(value: str) -> str:
    cleaned = "".join(char if char.isalnum() or char in {"-", "_"} else "_" for char in value.lower())
    return cleaned.strip("_") or "documento"


def _build_signed_text(document_text: str, signature: dict[str, Any], *, radicado: str | None = None) -> str:
    lines = [
        str(document_text or "").strip(),
        "",
        "",
        "Firma simple para radicacion",
        f"Nombre: {signature.get('full_name') or ''}",
        f"Documento: {signature.get('document_number') or ''}",
        f"Ciudad: {signature.get('city') or ''}",
        f"Fecha: {signature.get('date') or ''}",
    ]
    if radicado:
        lines.append(f"Radicado interno: {radicado}")
    return "\n".join(lines).strip() + "\n"


def _generate_docx_bytes(document_text: str, signature: dict[str, Any], *, radicado: str | None = None) -> bytes:
    if Document is None:
        return _build_signed_text(document_text, signature, radicado=radicado).encode("utf-8")

    document = Document()
    for paragraph_text in str(document_text or "").splitlines():
        document.add_paragraph(paragraph_text)
    document.add_paragraph("")
    document.add_paragraph("Firma simple para radicacion")
    document.add_paragraph(signature.get("full_name") or "")
    document.add_paragraph(f"Documento: {signature.get('document_number') or ''}")
    document.add_paragraph(f"{signature.get('city') or ''}, {signature.get('date') or ''}")
    if radicado:
        document.add_paragraph(f"Radicado interno: {radicado}")

    from io import BytesIO

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _escape_pdf_text(value: str) -> str:
    return value.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def _generate_pdf_bytes(document_text: str, signature: dict[str, Any], *, radicado: str | None = None) -> bytes:
    content_text = _build_signed_text(document_text, signature, radicado=radicado)
    lines = [line[:110] for line in content_text.splitlines()]
    y = 780
    operations = ["BT", "/F1 11 Tf", "50 780 Td"]
    first_line = True
    for line in lines:
        escaped = _escape_pdf_text(line)
        if not first_line:
            operations.append("0 -14 Td")
        operations.append(f"({escaped}) Tj")
        first_line = False
        y -= 14
        if y < 60:
            break
    operations.append("ET")
    stream = "\n".join(operations).encode("latin-1", errors="ignore")

    objects: list[bytes] = []
    objects.append(b"1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj\n")
    objects.append(b"2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj\n")
    objects.append(b"3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >> endobj\n")
    objects.append(b"4 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj\n")
    objects.append(f"5 0 obj << /Length {len(stream)} >> stream\n".encode("latin-1") + stream + b"\nendstream endobj\n")

    pdf = b"%PDF-1.4\n"
    offsets = [0]
    for obj in objects:
        offsets.append(len(pdf))
        pdf += obj
    xref_offset = len(pdf)
    pdf += f"xref\n0 {len(objects) + 1}\n".encode("latin-1")
    pdf += b"0000000000 65535 f \n"
    for offset in offsets[1:]:
        pdf += f"{offset:010d} 00000 n \n".encode("latin-1")
    pdf += f"trailer << /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF".encode("latin-1")
    return pdf


def create_signed_submission_artifacts(
    *,
    case_id: str,
    recommended_action: str,
    document_text: str,
    signature: dict[str, Any],
    radicado: str | None = None,
) -> dict[str, Any]:
    root = ensure_upload_root()
    target_dir = root / "submissions" / case_id
    target_dir.mkdir(parents=True, exist_ok=True)

    base_name = _safe_name(recommended_action or "documento")
    timestamp = datetime.utcnow().strftime("%Y%m%d%H%M%S")
    stem = f"{base_name}_{timestamp}"

    docx_path = target_dir / f"{stem}.docx"
    pdf_path = target_dir / f"{stem}.pdf"

    docx_bytes = _generate_docx_bytes(document_text, signature, radicado=radicado)
    pdf_bytes = _generate_pdf_bytes(document_text, signature, radicado=radicado)

    docx_path.write_bytes(docx_bytes)
    pdf_path.write_bytes(pdf_bytes)

    return {
        "docx_relative_path": docx_path.relative_to(root).as_posix(),
        "pdf_relative_path": pdf_path.relative_to(root).as_posix(),
        "docx_filename": docx_path.name,
        "pdf_filename": pdf_path.name,
        "signature": signature,
        "generated_at": datetime.utcnow().isoformat() + "Z",
    }
